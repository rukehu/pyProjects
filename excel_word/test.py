from docx import Document
from docx.shared import Inches
from docx import Document
from docx.enum.style import WD_STYLE_TYPE

document = Document()


idx = 0
styles = document.styles
for sty in styles:
    if sty.type == WD_STYLE_TYPE.PARAGRAPH:
        idx += 1
        idx_str = str(idx)
        document.add_paragraph(idx_str + sty.name, style=sty.name)


for idx in range(10):
    document.add_heading('Heading'+str(idx), level=idx)


document.save('demo.docx')