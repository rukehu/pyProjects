#!/usr/bin/python
# -*- coding: UTF-8 -*-

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from docx.shared import Inches
from docx.shared import RGBColor
import logging

logger = logging.getLogger('WordHandl')

FONT_NAME = 'Calibri'        # 设置默认字体
TAB_CLOUMN_CNT = 6           # 表格默认列数

class WrodHandl(object):
    def __init__(self):
        self._word_doc = None
        self._tab_cnt = 1

    def __get_tabshare_list(self, tab_shares):
        """
        获取table share reg list
        :param tab_shares:
        :return: list()
        """
        if len(tab_shares) > 0:
            return tab_shares.split()    # 以空字符为分割符

        return tab_shares

    def __get_tabhash_list(self, tab_hashs):
        if len(tab_hashs) > 0:
            return tab_hashs.splitlines()

        return tab_hashs


    def __get_tab_discriptions(self, disc):
        """
        获取寄存器描述信息列表
        :param disc:
        :return:list()
        """
        disc_list = list()
        str1 = 'Index Description'
        str2 = 'Memory Table Description'
        str3 = 'Description'
        lookup = False

        disc1 = None
        disc2 = None
        disc3 = None

        # 去掉多余的换行符
        swp_str = disc.splitlines()
        disc = str()
        for add_str in swp_str:
            _str = add_str.rstrip()
            disc += _str + ' '

        str_len = len(disc)
        idx = disc.find(str1)
        if idx >= 0:
            disc3 = disc[idx:str_len]
            str_len -= len(disc3)
            lookup = True

        if str_len > 0:
            swp_str = disc[0:str_len]
            idx = swp_str.find(str2)
            if idx >= 0:
                disc2 = swp_str[idx:str_len]
                str_len -= len(disc2)
                lookup = True

        if str_len > 0 or not lookup:
            swp_str = disc[0:str_len]
            idx = swp_str.find(str3)
            if idx >= 0:
                disc1 = swp_str[idx:str_len]
            else:
                disc1 = str3 + ': ' +swp_str

        if disc1 != None:
            disc_list.append(disc1)

        if disc2 != None:
            lower_str = str2
            rep = disc2.replace(str2, lower_str.lower().capitalize())
            disc_list.append(rep)

        if disc3 != None:
            lower_str = str1
            rep = disc3.replace(str1, lower_str.lower().capitalize())
            disc_list.append(rep)

        return disc_list

    def __write_hand_info(self, hand_info):
        """
        写入寄存器总体信息
        :param hand_info:
        :return:
        """
        have_key = False
        pgr = self._word_doc.add_heading(level=1)  # 写入寄存器表名
        pgr.paragraph_format.space_before = Pt(10)
        pgr.paragraph_format.space_after = Pt(10)
        run = pgr.add_run(hand_info['KEY_TYPE'])
        font = run.font
        font.name = FONT_NAME
        font.size = Pt(15)
        font.color.rgb = RGBColor(0x0, 0x0, 0x0)

        if 'TABLE_COUNT' in hand_info.keys():
            tab_cnt = 'Table count: ' + str(hand_info['TABLE_COUNT'])
            self._word_doc.add_paragraph(tab_cnt)

        tab_bitwidth = 'Table bit width: '
        if 'TABLE_BITS_WIDTH' in hand_info.keys():
            have_key = True
            tab_bitwidth += str(hand_info['TABLE_BITS_WIDTH'])
        elif 'TABLE_WIDTH' in hand_info.keys():
            tab_bitwidth += str(hand_info['TABLE_WIDTH'])
        if have_key:
            self._word_doc.add_paragraph(tab_bitwidth)

        if 'TABLE_TYPE' in hand_info.keys():
            tab_type = 'Table type: ' + hand_info['TABLE_TYPE'].lower().capitalize()
            self._word_doc.add_paragraph(tab_type)

        if 'HASH_ALGORITHM' in hand_info.keys():
            hash_list = self.__get_tabhash_list(hand_info['HASH_ALGORITHM'])
            self._word_doc.add_paragraph('Hash algorithm:')
            for hs in hash_list:
                self._word_doc.add_paragraph('.' + hs, style='List 2')

        if 'TABLE_SHARE' in hand_info.keys():
            share_list = self.__get_tabshare_list(hand_info['TABLE_SHARE'])
            self._word_doc.add_paragraph('Table share:')
            for share in share_list:
                self._word_doc.add_paragraph('.' + share, style='List 2')

        if 'entry_type' in hand_info.keys():
            tab_str = 'Entry type: ' + hand_info['entry_type']
            self._word_doc.add_paragraph(tab_str)

        if 'TABLE_DESCRIPTION' in hand_info.keys():
            desc_list = self.__get_tab_discriptions(hand_info['TABLE_DESCRIPTION'])
            for desc in desc_list:
                self._word_doc.add_paragraph(desc)


    def __write_table_reginfo(self, reg_info):
        """
        表格写入寄存器信息
        :param reg_tab:
        :return:
        """
        # 添加表格标题
        tab_name = 'Table ' + str(self._tab_cnt)
        _nam = ': ' + reg_info['KEY_TYPE']
        self._tab_cnt += 1
        tab_pgr = self._word_doc.add_paragraph(style='TOC Heading')
        tab_pgr.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        tab_run = tab_pgr.add_run(tab_name)

        font = tab_run.font
        font.name = FONT_NAME
        font.color.rgb = RGBColor(0x0, 0x0, 0x0)
        font.size = Pt(11)
        font.italic = True
        tab_run = tab_pgr.add_run(_nam)
        font = tab_run.font
        font.name = FONT_NAME
        font.color.rgb = RGBColor(0x0, 0x0, 0x0)
        font.size = Pt(11)
        font.italic = True
        # 设置段落距离
        tab_pgr.paragraph_format.space_before = Pt(10)
        tab_pgr.paragraph_format.space_after = Pt(5)

        bit_list = reg_info['bit_list']
        tab_row = len(bit_list) + 1  # 需添加表头
        head_arr = ['Bits', 'Name', 'Description', 'Default']

        have_note = False
        have_iskey = False

        if 'IS_KEY' in bit_list[0].keys():
            head_arr.append('IsKey')
            have_iskey = True

        for bit_info in bit_list:           # 查询是否存在note列
            if bit_info['NOTES'] != None and not have_note:
                head_arr.append('Notes')
                have_note = True
                break

        # 设置表格列数与宽度
        tab_cloumn = TAB_CLOUMN_CNT
        width_0 = Inches(0.7)
        width_1 = Inches(1.6)
        width_2 = Inches(1.9)
        width_3 = Inches(0.7)
        width_4 = Inches(0.5)
        width_5 = Inches(0.8)

        iskey_idx = 4
        note_idx = 5
        if not have_iskey and not have_note:
            tab_cloumn -= 2
            width_1 = Inches(2.0)
            width_2 = Inches(2.8)
            width_3 = Inches(0.7)

        elif have_iskey and not have_note:
            tab_cloumn -= 1
            width_1 = Inches(1.8)
            width_2 = Inches(2.5)
            width_3 = Inches(0.7)
            width_4 = Inches(0.5)

        elif have_note and not have_iskey:
            tab_cloumn -= 1
            note_idx = 4
            width_0 = Inches(0.7)
            width_1 = Inches(1.8)
            width_2 = Inches(2.2)
            width_3 = Inches(0.7)
            width_4 = Inches(0.8)

        reg_tab = self._word_doc.add_table(tab_row, tab_cloumn, 'Table Grid')

        # 设置表格宽度

        reg_tab.autofit = False
        for idx in range(tab_row):
            reg_tab.cell(idx, 0).width = width_0
            reg_tab.cell(idx, 1).width = width_1
            reg_tab.cell(idx, 2).width = width_2
            reg_tab.cell(idx, 3).width = width_3
            if have_note and have_iskey:
                reg_tab.cell(idx, iskey_idx).width = width_4
                reg_tab.cell(idx, note_idx).width = width_5
            elif have_iskey:
                reg_tab.cell(idx, iskey_idx).width = width_4
            elif have_note:
                reg_tab.cell(idx, note_idx).width = width_4

        for idx in range(tab_cloumn):
            tab_run = reg_tab.cell(0, idx).paragraphs[0].add_run(head_arr[idx])
            tab_run.bold = True
            tab_run.font.italic = True

        row_idx = 1
        for bit_info in bit_list:
            tab_cells = reg_tab.rows[row_idx].cells

            bits = bit_info['FIELD_BITS']
            if bits == None:
                bits = bit_info['SUB_FIELD_BITS']
            tab_cells[0].text = str(bits)

            name = bit_info['FIELD_NAME']
            if name == None:
                name = bit_info['SUB_FIELD_NAME']
            tab_cells[1].text = name

            descrip = bit_info['DESCRIPTION']
            if descrip != None:
                tab_cells[2].text = descrip
            default = bit_info['DEFAULT_VALUE']
            if default != None:
                tab_cells[3].text = str(default)

            if have_iskey:
                iskey = bit_info['IS_KEY']
                if iskey != None:
                    tab_cells[iskey_idx].text = str(iskey)

            note = bit_info['NOTES']
            if have_note and note != None:
                tab_cells[note_idx].text = str(note)

            row_idx += 1

    def set_table_number(self, base_num):
        if base_num > 0:
            self._tab_cnt = base_num

    def open_word(self, word_path):
        """
        打开一个word文档，不存在则创建
        :param word_path:
        :return:
        """
        try:
            self._word_doc = Document(word_path)
        except:
            logger.info('Can not open this file:', word_path)
            self._word_doc = None
            return

        self._word_doc.styles['Normal'].font.name = FONT_NAME

    def write_register_toword(self, reg_name, reg_info):
        """
        将给与寄存器信息写入word文档
        :param reg_info:
        :return:
        """
        reg_head = dict()

        if self._word_doc == None:
            logger.info('The document is not open.')
            return

        key_list = reg_info['head_info'].keys()
        for _key in key_list:
            if reg_info['head_info'][_key] != None:
                reg_head[_key] = reg_info['head_info'][_key]

        reg_head['KEY_TYPE'] = reg_name
        self.__write_hand_info(reg_head)

        # 将寄存器信息写入word表格
        reg_list = reg_info['reg_list']
        for reg_tab in reg_list:
            self.__write_table_reginfo(reg_tab)

    def write_word_end(self, word_url):
        # self._word_doc.add_page_break()
        self._word_doc.save(word_url)
